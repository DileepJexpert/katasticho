from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PRIVATE_LABEL_FMCG_FIELD_SALES_QA_RUNBOOK.md"
OUTPUT = ROOT / "docs" / "PRIVATE_LABEL_FMCG_FIELD_SALES_QA_RUNBOOK.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TABLE_FILL = "E8EEF5"
MUTED = "5B6573"
TABLE_WIDTH = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Calibri", size=11, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clean_inline(text):
    return re.sub(r"\*\*|`", "", text).strip()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_table(doc, lines):
    rows = []
    for line in lines:
        cells = [clean_inline(c) for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r"\s*:?-+:?\s*", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    cols = len(rows[0])
    if cols == 2:
        widths = [2200, 7160]
    elif cols == 3:
        widths = [1700, 2500, 5160]
    elif cols == 4:
        widths = [1700, 2500, 2400, 2760]
    elif cols == 5:
        widths = [1650, 2050, 1750, 1800, 2110]
    else:
        widths = [TABLE_WIDTH // cols] * cols
        widths[-1] += TABLE_WIDTH - sum(widths)

    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, row_data in enumerate(rows):
        row = table.rows[row_index]
        for col_index in range(cols):
            value = row_data[col_index] if col_index < len(row_data) else ""
            cell = row.cells[col_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(value)
            set_font(run, size=9.5, bold=row_index == 0, color=DARK_BLUE if row_index == 0 else None)
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("KATASTICHO ERP  |  PRIVATE-LABEL FMCG QA")
    set_font(run, size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, table_lines)
            continue

        match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if match:
            level = len(match.group(1))
            text = clean_inline(match.group(2))
            if first_heading and level == 1:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                run = paragraph.add_run(text)
                set_font(run, size=24, color="0B2545", bold=True)
                first_heading = False
            else:
                doc.add_heading(text, level=level)
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", line)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(clean_inline(bullet.group(1)))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.add_run(clean_inline(numbered.group(1)))
            index += 1
            continue

        paragraph = doc.add_paragraph()
        if line.startswith("**") and line.endswith("**"):
            run = paragraph.add_run(clean_inline(line))
            set_font(run, bold=True, color=DARK_BLUE)
        else:
            paragraph.add_run(clean_inline(line))
        index += 1

    props = doc.core_properties
    props.title = "Private-Label FMCG Distributor and Field Sales QA Runbook"
    props.subject = "Repeatable manual regression testing flow for Katasticho ERP"
    props.author = "Katasticho ERP"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
